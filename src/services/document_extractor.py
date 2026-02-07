"""
Document Extractor - parses attachments from Fedresurs messages.
Extracts structured data from PDF/DOCX files (ЕГРН extracts, appraisal reports).
"""
import re
import logging
from typing import Optional, Dict, Any, List
from pathlib import Path
import tempfile

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extracts structured data from property documents."""

    # Regex patterns for common data
    CADASTRAL_PATTERN = r'\d{2}:\d{2}:\d{6,7}:\d{1,}'
    AREA_PATTERN = r'(?:площадь|площадью)[:\s]*(\d+[,.]?\d*)\s*(?:кв\.?\s*м|м²)'
    INN_PATTERN = r'\b\d{10}\b|\b\d{12}\b'

    def __init__(self):
        """Initialize document extractor."""
        self.supported_formats = []

        if PyPDF2 or pdfplumber:
            self.supported_formats.append("pdf")
        if DocxDocument:
            self.supported_formats.append("docx")

        logger.info(f"DocumentExtractor initialized. Supported formats: {self.supported_formats}")

    async def extract_from_file(self, file_path: str, document_type: str = "unknown") -> Dict[str, Any]:
        """
        Extract data from document file.

        Args:
            file_path: Path to the document
            document_type: Type hint (egr_extract, appraisal_report, etc.)

        Returns:
            Extracted structured data with metadata
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            logger.error(f"File not found: {file_path}")
            return {"error": "file_not_found"}

        extension = file_path_obj.suffix.lower().lstrip('.')

        if extension not in self.supported_formats:
            logger.warning(f"Unsupported file format: {extension}")
            return {"error": f"unsupported_format: {extension}"}

        try:
            if extension == "pdf":
                return await self._extract_from_pdf(file_path, document_type)
            elif extension == "docx":
                return await self._extract_from_docx(file_path, document_type)
            else:
                return {"error": "unsupported_format"}

        except Exception as e:
            logger.error(f"Failed to extract data from {file_path}: {e}")
            return {"error": str(e)}

    async def _extract_from_pdf(self, file_path: str, document_type: str) -> Dict[str, Any]:
        """Extract data from PDF file."""
        extracted = {
            "file_type": "pdf",
            "document_type": document_type,
            "text": "",
            "cadastral_numbers": [],
            "areas": [],
            "inns": [],
            "encumbrances": [],
            "metadata": {}
        }

        # Try pdfplumber first (better for tables)
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            extracted["text"] += text + "\n"
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")

        # Fallback to PyPDF2
        if not extracted["text"] and PyPDF2:
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        extracted["text"] += page.extract_text() + "\n"
            except Exception as e:
                logger.error(f"PyPDF2 extraction failed: {e}")

        # Parse extracted text
        if extracted["text"]:
            self._parse_cadastral_numbers(extracted)
            self._parse_areas(extracted)
            self._parse_inns(extracted)

            # Special parsing for ЕГРН extracts
            if "егрн" in document_type.lower() or "выписк" in extracted["text"].lower():
                self._parse_egr_extract(extracted)

            # Parse appraisal reports
            if "оценк" in extracted["text"].lower():
                self._parse_appraisal_report(extracted)

        return extracted

    async def _extract_from_docx(self, file_path: str, document_type: str) -> Dict[str, Any]:
        """Extract data from DOCX file."""
        if not DocxDocument:
            return {"error": "python-docx not installed"}

        extracted = {
            "file_type": "docx",
            "document_type": document_type,
            "text": "",
            "cadastral_numbers": [],
            "areas": [],
            "inns": [],
            "encumbrances": [],
            "metadata": {}
        }

        try:
            doc = DocxDocument(file_path)

            # Extract all paragraphs
            for para in doc.paragraphs:
                extracted["text"] += para.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted["text"] += cell.text + " "
                extracted["text"] += "\n"

            # Parse extracted text
            self._parse_cadastral_numbers(extracted)
            self._parse_areas(extracted)
            self._parse_inns(extracted)

            if "егрн" in document_type.lower():
                self._parse_egr_extract(extracted)

            if "оценк" in extracted["text"].lower():
                self._parse_appraisal_report(extracted)

        except Exception as e:
            logger.error(f"Failed to extract from DOCX: {e}")
            extracted["error"] = str(e)

        return extracted

    def _parse_cadastral_numbers(self, extracted: Dict):
        """Extract cadastral numbers from text."""
        matches = re.findall(self.CADASTRAL_PATTERN, extracted["text"])
        extracted["cadastral_numbers"] = list(set(matches))

    def _parse_areas(self, extracted: Dict):
        """Extract area measurements from text."""
        matches = re.findall(self.AREA_PATTERN, extracted["text"], re.IGNORECASE)
        areas = []
        for match in matches:
            try:
                area_value = float(match.replace(',', '.'))
                areas.append(area_value)
            except ValueError:
                pass
        extracted["areas"] = list(set(areas))

    def _parse_inns(self, extracted: Dict):
        """Extract INN numbers from text."""
        matches = re.findall(self.INN_PATTERN, extracted["text"])
        extracted["inns"] = list(set(matches))

    def _parse_egr_extract(self, extracted: Dict):
        """Parse specific data from ЕГРН extract."""
        text_lower = extracted["text"].lower()

        # Find encumbrances section
        encumbrance_markers = ["обременения", "ограничения", "залог", "ипотека", "аренда"]
        for marker in encumbrance_markers:
            if marker in text_lower:
                # Simple extraction: get text around marker
                start = text_lower.find(marker)
                snippet = extracted["text"][start:start+500]

                # Look for common encumbrance types
                if "ипотека" in snippet.lower():
                    extracted["encumbrances"].append({
                        "type": "mortgage",
                        "snippet": snippet[:200]
                    })
                if "аренда" in snippet.lower():
                    extracted["encumbrances"].append({
                        "type": "lease",
                        "snippet": snippet[:200]
                    })
                if "арест" in snippet.lower():
                    extracted["encumbrances"].append({
                        "type": "arrest",
                        "snippet": snippet[:200]
                    })

        # Find cadastral value
        cadastral_value_pattern = r'кадастровая\s+стоимость[:\s]*(\d+[\s,.]?\d*)'
        match = re.search(cadastral_value_pattern, text_lower)
        if match:
            try:
                value = match.group(1).replace(' ', '').replace(',', '.')
                extracted["metadata"]["cadastral_value"] = float(value)
            except ValueError:
                pass

    def _parse_appraisal_report(self, extracted: Dict):
        """Parse appraisal report data."""
        text = extracted["text"]

        # Find market value
        market_value_patterns = [
            r'рыночная\s+стоимость[:\s]*(\d+[\s,.]?\d*)',
            r'итоговая\s+величина\s+стоимости[:\s]*(\d+[\s,.]?\d*)'
        ]

        for pattern in market_value_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    value = match.group(1).replace(' ', '').replace(',', '.')
                    extracted["metadata"]["market_value"] = float(value)
                    break
                except ValueError:
                    pass

        # Find appraisal date
        date_pattern = r'дата\s+оценки[:\s]*(\d{2}[.\-/]\d{2}[.\-/]\d{4})'
        match = re.search(date_pattern, text, re.IGNORECASE)
        if match:
            extracted["metadata"]["appraisal_date"] = match.group(1)

    async def extract_from_attachment(
        self,
        attachment_data: bytes,
        filename: str,
        document_type: str = "unknown"
    ) -> Dict[str, Any]:
        """
        Extract data from attachment bytes.

        Args:
            attachment_data: File bytes
            filename: Original filename
            document_type: Document type hint

        Returns:
            Extracted data
        """
        # Save to temp file
        suffix = Path(filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(attachment_data)
            tmp_path = tmp.name

        try:
            result = await self.extract_from_file(tmp_path, document_type)
            result["original_filename"] = filename
            return result
        finally:
            # Clean up temp file
            try:
                Path(tmp_path).unlink()
            except Exception as e:
                logger.warning(f"Failed to delete temp file {tmp_path}: {e}")

    def detect_document_type(self, filename: str, text_snippet: str = "") -> str:
        """
        Detect document type from filename and content.

        Args:
            filename: File name
            text_snippet: Optional text snippet from document

        Returns:
            Document type identifier
        """
        filename_lower = filename.lower()
        text_lower = text_snippet.lower()

        # ЕГРН extracts
        if any(word in filename_lower for word in ["егрн", "выписка", "extract"]):
            return "egr_extract"

        if "выписк" in text_lower and "росреестр" in text_lower:
            return "egr_extract"

        # Appraisal reports
        if any(word in filename_lower for word in ["оценка", "отчет", "appraisal"]):
            return "appraisal_report"

        if "рыночная стоимость" in text_lower or "оценщик" in text_lower:
            return "appraisal_report"

        # Default
        return "unknown"
